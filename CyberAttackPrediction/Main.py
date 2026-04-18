# ==============================================================================
# FILE: Main.py
# PROJECT: CyberShield AI — Cyber Attack Prediction System (2026 Stable)
# TECH STACK: Python 3.13 | Flask 3.1.3 | Pandas 3.0.2 | Scikit-Learn 1.8.0
# ==============================================================================
#
# WHAT THIS FILE IS:
#   This is the "brain and spine" of the entire web application.
#   It does three main jobs:
#     1. RUNS THE WEBSITE — handles every page (Home, Login, Predict, Train, etc.)
#     2. HANDLES SECURITY — manages user accounts, login sessions, and CSRF protection
#     3. RUNS THE AI — loads the trained model and uses it to classify network attacks
#
# HOW IT WORKS (Simple Version):
#   When a user visits the website in their browser, Flask (a web server library)
#   receives the request, runs the matching Python function, and sends back an
#   HTML page. Think of it like a waiter — it takes the order (request), goes to
#   the kitchen (processes data), and brings back the food (webpage).
#
# KEY SECTIONS INSIDE THIS FILE:
#   • IMPORTS        — loads all the tools/libraries this script needs
#   • SECURITY SETUP — CSRF protection and response headers to keep users safe
#   • MODEL LOADER   — loads the pre-trained AI from disk at startup (so it's fast)
#   • USER SYSTEM    — signup, login, logout, password management, admin controls
#   • PREDICT        — takes a user's uploaded CSV, runs the AI, returns results
#   • TRAIN          — triggers the training script and reloads the new model
#   • EXPLORER       — lets logged-in users browse and read project files in-browser
#   • ERROR HANDLERS — friendly 404 page and catch-all error responses
#   • STARTUP        — runs the web server and auto-opens the browser when ready
#
# HOW TO RUN:
#   Just use the launcher (launcher.py → Option 1), or run:
#     python Main.py
#   The site will start at http://127.0.0.1:2026/
#================================================================================
import pandas as pd
import numpy as np
import joblib
import os
import webbrowser
import json
from threading import Timer
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.model_selection import train_test_split
from sklearn.ensemble import StackingClassifier
from sklearn.ensemble import RandomForestClassifier
from sklearn.tree import DecisionTreeClassifier
from sklearn.linear_model import LogisticRegression
from sklearn.neural_network import MLPClassifier
import secrets
import subprocess
import sys

#=================flask code starts here
from flask import Flask, render_template, request, redirect, url_for, session, flash, send_from_directory, jsonify, render_template_string
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from train_model import run_training
import markdown2
from pygments import highlight
from pygments.lexers import get_lexer_for_filename
from pygments.formatters import HtmlFormatter
import nbformat
from nbconvert import HTMLExporter
import logging
from docx import Document
from docx.shared import Pt

# Silence repetitive Heartbeat logs for a cleaner terminal experience
class HeartbeatFilter(logging.Filter):
    def filter(self, record):
        # Exclude any log record containing the heartbeat URI
        return "/api/heartbeat" not in record.getMessage()

log = logging.getLogger('werkzeug')
log.addFilter(HeartbeatFilter())

# Load environment variables
load_dotenv()

app = Flask(__name__, static_url_path='')
# Use a derived secret key for session integrity; fallback to a generated secret if environment is missing
app.secret_key = os.getenv('FLASK_SECRET_KEY', secrets.token_hex(32) if not os.getenv('FLASK_SECRET_KEY') else os.getenv('FLASK_SECRET_KEY'))

# PROPRIETARY ADMIN SECURITY CONFIG (Loaded from Environment)
ADMIN_ID = os.getenv('ADMIN_USER', 'admin')
ADMIN_HASH = os.getenv('ADMIN_HASH', ADMIN_HASH if 'ADMIN_HASH' in locals() else "scrypt:32768:8:1$Cw2lrwbxEaJWsMvM$7d647e7e6e63deab9540306b28ceb1ea622ce0632d2d1b4e10807b507efe42355fab976a9acf2b5c0bc8e4a07ed75c2ba5f8e19330c7e3aac45ee82ceab4ceda")

ALLOWED_EXTENSIONS = {'csv'}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Global variables to store loaded models/scalers
rf_model = None
scaler = None
labels = None
label_encoder = None
feature_columns = None
system_status = "ready"  # Track if training is in progress
PULSE_DETECTED = False # Pulse Detection Flag
browser_timer = None   # Global Browser Launch Timer

# --- Global Security Middlewares ---
@app.context_processor
def inject_csrf_token():
    return dict(csrf_token=lambda: session.get('csrf_token'))

@app.before_request
def security_pre_check():
    """Generates CSRF tokens and enforces security validation on sensitive routes."""
    # 1. Ensure CSRF token exists for the session
    if 'csrf_token' not in session:
        session['csrf_token'] = secrets.token_hex(32)

    # 2. CSRF Validation for all state-changing requests (POST/PUT/DELETE)
    if request.method in ['POST', 'PUT', 'DELETE']:
        # Exempt specific public routes if necessary (None currently)
        exempt_routes = ['/UserLoginAction', '/SignupAction'] # Optional: allow initial auth without token if needed
        if request.path in exempt_routes:
             return
             
        # Check header or form data for the token
        token = request.headers.get('X-CSRF-Token') or request.form.get('csrf_token')
        if not token or token != session.get('csrf_token'):
            return jsonify({"status": "error", "message": "Security Violation: CSRF Token Invalid or Missing"}), 403

@app.after_request
def apply_optimization_headers(response):
    """Hardens the response headers and enables ETag-based caching for optimal reload performance."""
    # 1. Standard security headers
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "SAMEORIGIN"

    # 2. Performance: Enable ETag generation for all successful GET requests
    # This allows the browser to perform 'Conditional GETs' (304 Not Modified)
    if request.method == 'GET' and response.status_code == 200:
        response.add_etag()
        response.make_conditional(request)

    # 3. Cache-Control Strategy:
    # - Sensitive user data routes: Strict no-store
    # - Public Templates: allow browser to check for changes (must-revalidate)
    if 'user' in session:
        # For logged-in users, we still need to prevent stale data leaks
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    else:
        # Public landing/login/signup pages: allow caching with revalidation (fastest feel)
        response.headers["Cache-Control"] = "public, no-cache, must-revalidate"
    
    return response

def load_ml_model():
    """Loads the pre-trained model and preprocessing tools from disk with integrity checks."""
    global rf_model, scaler, labels, label_encoder, feature_columns
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        model_path = os.path.join(base_dir, "model", "trained_rf_model.pkl")
        
        if os.path.exists(model_path):
            # Basic integrity check: File size must be reasonable
            if os.path.getsize(model_path) < 1000: # Less than 1KB is likely corruption
                 print(f"[Security Check] Model file {model_path} appears corrupted or too small.")
                 return False

            model_data = joblib.load(model_path)
            # Verify structure of the loaded object
            required_keys = ['rf_model', 'scaler', 'labels', 'label_encoder', 'feature_columns']
            if not all(k in model_data for k in required_keys):
                print("[Security Check] Model data integrity compromised: Missing required components.")
                return False

            rf_model = model_data['rf_model']
            scaler = model_data['scaler']
            labels = model_data['labels']
            label_encoder = model_data['label_encoder']
            feature_columns = model_data['feature_columns']
            print(f"Successfully loaded pre-trained model from {model_path}.")
            return True
        else:
            print(f"Pre-trained model not found at {model_path}. Please run train_model.py first.")
            return False
    except Exception as e:
        print(f"Error loading model: {e}")
        return False

# Load model at startup
load_ml_model()

@app.route('/')
@app.route('/index')
def index():
    # Allow manual reset via query param or button
    if request.args.get('reset') == 'true':
        session.pop('last_result', None)
        session.pop('last_page_type', None)

    # Home page always shows the landing page.
    # Previous results are only restored on the /Predict page, not here.
    return render_template('index.html')

@app.route('/HowItWorks')
def HowItWorks():
    return render_template('HowItWorks.html')

@app.route('/Predict')
@app.route('/PredictView')
def predictView():
    if 'user' not in session:
        return redirect(url_for('UserLogin'))
    
    # Check if model is actually ready before allowing prediction
    base_dir = os.path.dirname(os.path.abspath(__file__))
    model_path = os.path.join(base_dir, "model", "trained_rf_model.pkl")
    if not os.path.exists(model_path):
        flash("Model not found. Please train the model before running predictions.", "warning")
        return redirect(url_for('train_view'))
    
    # If reset is requested, clear the last result
    if request.args.get('reset') == 'true':
        session.pop('last_result', None)
        session.pop('last_page_type', None)
        return redirect(url_for('predictView'))
        
    # Allow the UI to know if a result already exists in this session
    last_result = session.get('last_result')
    return render_template('Predict.html', page_type='input', last_result=last_result)

@app.route('/api/clear-session', methods=['POST'])
def clear_session():
    """API endpoint to clear all app results and terminate session for system reset."""
    if not session.get('user'):
        return jsonify({"status": "error", "message": "Unauthorized"}), 401
        
    # Clear all state results (but NOT filesystem credentials as requested)
    session.pop('last_result', None)
    session.pop('train_result', None)
    session.pop('last_page_type', None)
    session.pop('last_feature_columns', None)
    
    # Also log the user out as requested for a total state reset
    session.pop('user', None)
    session.pop('is_ultimate', None)
    
    return jsonify({"status": "success", "message": "System data cleared and session terminated."})

@app.route('/ClearResults', methods=['POST'])
def ClearResults():
    """Wipes the last analysis results and resets the dashboard state for a fresh start."""
    if 'user' not in session:
        return jsonify({"status": "error", "message": "Unauthorized"}), 401
        
    session.pop('last_result', None)
    session.pop('last_page_type', None)
    session.pop('last_feature_columns', None)
    
    return jsonify({"status": "success", "message": "Analysis results wiped successfully."})

@app.route('/api/heartbeat', methods=['GET', 'POST'])
def heartbeat():
    """Silent Heartbeat for Smart Launch and System Status Sync."""
    global browser_timer, PULSE_DETECTED
    
    # Pulse detected! Mark system as having an active session
    PULSE_DETECTED = True
    
    # If a heartbeat arrives, an active tab exists - cancel the startup browser launch
    if browser_timer and browser_timer.is_alive():
        print("[Pulse Detection] Pulse detected from existing tab. Cancelling auto-launch.")
        browser_timer.cancel()
        
    return jsonify({"status": system_status, "health": "online", "pulse": "active"})

# --- Security & User Management ---
USERS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "users.json")

def load_users():
    """Enterprise-grade user directory loader with auto-initialization for master admin."""
    if not os.path.exists(USERS_FILE):
        users = {
            ADMIN_ID: {
                "username": ADMIN_ID,
                "password": ADMIN_HASH,
                "role": "admin"
            }
        }
        save_users(users)
        return users
    
    try:
        with open(USERS_FILE, 'r') as f:
            users = json.load(f)
            # Self-healing: Restore master admin if missing from archives
            if ADMIN_ID not in users:
                users[ADMIN_ID] = {
                    "username": ADMIN_ID,
                    "password": ADMIN_HASH,
                    "role": "admin"
                }
                save_users(users)
            return users
    except Exception as e:
        print(f"[Identity Cache] Error: {e}")
        return {}

def save_users(users):
    with open(USERS_FILE, 'w') as f:
        json.dump(users, f, indent=4)

@app.route('/UserLogin', methods=['GET', 'POST'])
def UserLogin():
    if 'user' in session:
        return redirect(url_for('train_view'))
    base_dir = os.path.dirname(os.path.abspath(__file__))
    creds_file = os.path.join(base_dir, "saved_creds.json")
    has_saved = os.path.exists(creds_file)
    return render_template('UserLogin.html', has_saved=has_saved)

@app.route('/VerifySavedLogin', methods=['POST'])
def VerifySavedLogin():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    creds_file = os.path.join(base_dir, "saved_creds.json")
    
    if os.path.exists(creds_file):
        try:
            with open(creds_file, 'r') as f:
                creds = json.load(f)
            
            username = creds.get('user')
            signature = creds.get('token')
            
            if not username or not signature:
                return jsonify({"status": "error", "message": "Corrupted credentials payload"})

            # Generate expected signature to verify identity without password storage
            import hmac, hashlib
            expected = hmac.new(app.secret_key.encode(), username.encode(), hashlib.sha256).hexdigest()
            
            if hmac.compare_digest(signature, expected):
                # Identity verified via cryptographic signature
                session['user'] = username
                if username == ADMIN_ID:
                    session['is_ultimate'] = True
                session['remembered'] = True
                return jsonify({"status": "success"})
                
        except Exception as e:
            print(f"[Security Layer] Credential Verification Error: {e}")
            
    return jsonify({"status": "error", "message": "Saved credentials invalid or signature mismatch"})

@app.route('/ResetCredentials', methods=['POST'])
def ResetCredentials():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    creds_file = os.path.join(base_dir, "saved_creds.json")
    if os.path.exists(creds_file):
        os.remove(creds_file)
    return jsonify({"status": "success"})

@app.route('/CheckUsername', methods=['POST'])
def CheckUsername():
    """Real-time identity verification for signup process."""
    data = request.get_json()
    username = data.get('username', '').strip()
    
    if not username:
        return jsonify({"status": "available", "valid": False})
        
    users = load_users()
    if username in users or username.lower() == "admin":
        return jsonify({"status": "taken", "message": "Identity ID already registered in Archives!"})
    
    return jsonify({"status": "available", "message": "Identity available for registration."})

@app.route('/UserLoginAction', methods=['POST'])
def UserLoginAction():
    username = request.form.get('t1')
    password = request.form.get('t2')
    remember = request.form.get('remember') == 'on'
    
    # SECURE ADMIN BYPASS CHECK (Hashed)
    if username == ADMIN_ID and check_password_hash(ADMIN_HASH, password):
        session['user'] = ADMIN_ID
        session['is_ultimate'] = True
        if remember:
            # SECURE PERSISTENCE: Store a cryptographic signature, never the plain password
            import hmac, hashlib
            token = hmac.new(app.secret_key.encode(), ADMIN_ID.encode(), hashlib.sha256).hexdigest()
            
            base_dir = os.path.dirname(os.path.abspath(__file__))
            with open(os.path.join(base_dir, "saved_creds.json"), 'w') as f:
                json.dump({"user": ADMIN_ID, "token": token}, f)
            session['remembered'] = True
        return jsonify({"status": "success", "message": "Ultimate Admin access granted!", "redirect": url_for('train_view')})

    users = load_users()
    user_data = users.get(username)
    
    if not user_data:
        return jsonify({"status": "error", "error_type": "invalid_username", "message": "Identification failed: Username not found."}), 401

    if check_password_hash(user_data['password'], password):
        session['user'] = user_data['username']
        
        # Save credentials locally if "Remember Me" is checked
        if remember:
            # SECURE PERSISTENCE: Store a cryptographic signature, never the plain password
            import hmac, hashlib
            token = hmac.new(app.secret_key.encode(), username.encode(), hashlib.sha256).hexdigest()
            
            base_dir = os.path.dirname(os.path.abspath(__file__))
            with open(os.path.join(base_dir, "saved_creds.json"), 'w') as f:
                json.dump({"user": username, "token": token}, f)
            session['remembered'] = True
        
        return jsonify({"status": "success", "message": f"Welcome, {username}!", "redirect": url_for('train_view')})
    else:
        return jsonify({"status": "error", "error_type": "invalid_password", "message": "Identification failed: Incorrect password key."}), 401

@app.route('/ResetPasswordAction', methods=['POST'])
def ResetPasswordAction():
    data = request.get_json()
    username = data.get('username')
    new_password = data.get('new_password')
    auto_login = data.get('auto_login', False)
    
    users = load_users()
    if username not in users:
        return jsonify({"status": "error", "message": "User not found in Archives"}), 404
        
    users[username]['password'] = generate_password_hash(new_password)
    save_users(users)
    
    if auto_login:
        session['user'] = username
        return jsonify({"status": "success", "message": "Identity restored. Synchronizing session...", "redirect": url_for('train_view')})
    
    return jsonify({"status": "success", "message": "Security keys updated. Proceed to manual login."})

@app.route('/Signup')
def Signup():
    if 'user' in session:
        return redirect(url_for('train_view'))
    return render_template('Signup.html')

@app.route('/SignupAction', methods=['POST'])
def SignupAction():
    username = request.form.get('username')
    password = request.form.get('password')
    
    users = load_users()
    # Check for both existing users and the reserved master admin identity
    if username in users or username.lower() == "admin":
        return jsonify({"status": "error", "message": "Identity ID already registered in Archives!"}), 400
        
    users[username] = {
        "username": username,
        "password": generate_password_hash(password),
        "role": "user"
    }
    save_users(users)
    return jsonify({
        "status": "success", 
        "message": "Account created successfully!",
        "user_data": {"username": username, "password": password}
    })

@app.route('/Account')
def Account():
    if 'user' not in session:
        return redirect(url_for('UserLogin'))
    users = load_users()
    user_data = users.get(session['user'])
    return render_template('AccountSettings.html', user_data=user_data)

@app.route('/UpdateAccountAction', methods=['POST'])
def UpdateAccountAction():
    if 'user' not in session:
        return redirect(url_for('UserLogin'))
        
    new_username = request.form.get('username')
    new_password = request.form.get('password')
    old_username = session['user']
    
    users = load_users()
    if old_username not in users:
        # Check for ultimate admin bypass
        if old_username == "admin" and session.get('is_ultimate'):
            # Ultimate admin can update themselves too (stored in file if they want)
            user_data = {"username": "admin", "password": generate_password_hash("admin"), "role": "admin"}
        else:
            flash("User session expired. Please login again.", "danger")
            return redirect(url_for('UserLogin'))
    else:
        user_data = users.pop(old_username)
    
    # Update username if changed and not taken
    if new_username != old_username:
        if new_username in users:
            users[old_username] = user_data # Restore
            flash("New username already taken!", "danger")
            return redirect(url_for('Account'))
        user_data['username'] = new_username
        session['user'] = new_username
        
    # Update password if provided
    if new_password:
        user_data['password'] = generate_password_hash(new_password)
        
    users[user_data['username']] = user_data
    save_users(users)
    flash("Profile updated successfully!", "success")
    return redirect(url_for('Account'))

@app.route('/AdminGetUsers')
def AdminGetUsers():
    """Returns a list of all users for the dashboard."""
    if 'user' not in session or (session.get('user') != 'admin' and not session.get('is_ultimate')):
        return jsonify({"status": "error", "message": "Unauthorized"}), 403
    
    users = load_users()
    user_list = []
    for uname, data in users.items():
        user_list.append({
            "username": uname,
            "role": data.get('role', 'user')
        })
    return jsonify(user_list)

@app.route('/AdminDeleteUser', methods=['POST'])
def AdminDeleteUser():
    """Deletes a user account (Admin only)."""
    if 'user' not in session or (session.get('user') != 'admin' and not session.get('is_ultimate')):
        return jsonify({"status": "error", "message": "Unauthorized"}), 403
    
    target_user = request.json.get('username')
    if target_user == "admin" or target_user == session['user']:
        return jsonify({"status": "error", "message": "Cannot delete root admin or current session"}), 400
        
    users = load_users()
    if target_user in users:
        del users[target_user]
        save_users(users)
        return jsonify({"status": "success", "message": f"User {target_user} deleted"})
    return jsonify({"status": "error", "message": "User not found"}), 404

@app.route('/AdminResetPassword', methods=['POST'])
def AdminResetPassword():
    """Allows an administrator to reset a user's password."""
    if 'user' not in session or (session.get('user') != 'admin' and not session.get('is_ultimate')):
        return jsonify({"status": "error", "message": "Unauthorized"}), 403
    
    data = request.json
    target_user = data.get('username')
    new_password = data.get('password')
    
    if not target_user or not new_password:
        return jsonify({"status": "error", "message": "Username and new password are required"}), 400
        
    if target_user == "admin" and not session.get('is_ultimate'):
         return jsonify({"status": "error", "message": "Only the developer can reset the root account"}), 400

    users = load_users()
    if target_user in users:
        users[target_user]['password'] = generate_password_hash(new_password)
        save_users(users)
        return jsonify({"status": "success", "message": f"Credentials for {target_user} have been updated"})
    
    return jsonify({"status": "error", "message": "User not found"}), 404

@app.route('/Logout')
def Logout():
    """Securely terminates the user session and clears all cached results."""
    # Preserve only the pulse state if applicable
    pulse_state = session.get('pulse_detected', False)
    session.clear()
    session['pulse_detected'] = pulse_state
    
    flash("Neural session terminated: You have been safely logged out.", "info")
    return redirect(url_for('index'))

@app.route('/PredictAction', methods=['POST'])
def PredictAction():
    if 'user' not in session:
        return redirect(url_for('UserLogin'))
        
    global rf_model, scaler, labels, label_encoder
    
    if rf_model is None:
        if not load_ml_model():
            flash("Model not ready. Please run the training process first.", "warning")
            return redirect(url_for('train_view'))

    try:
        # Load test data from the uploaded file
        if 't1' not in request.files:
            flash("No file was uploaded. Please select a CSV file.", "error")
            return redirect(url_for('predictView'))
            
        file = request.files['t1']
        if file.filename == '':
            flash("No file was selected. Please try again.", "error")
            return redirect(url_for('predictView'))
            
        if not allowed_file(file.filename):
            flash("Invalid file type. Only CSV datasets are permitted.", "error")
            return redirect(url_for('predictView'))
            
        filename = secure_filename(file.filename)
        upload_path = os.path.join("Dataset", "uploaded_" + filename)
        file.save(upload_path)
        testData_df = pd.read_csv(upload_path)
        
        # SCHEMA VALIDATION: Ensure uploaded dataset has the expected features
        if feature_columns is not None:
            missing_cols = [c for c in feature_columns if c not in testData_df.columns]
            if missing_cols:
                flash(f"Invalid Prediction Data: Missing features {missing_cols}", "error")
                return redirect(url_for('predictView'))
            
            # Reorder columns to match the model's expected input order
            testData_df = testData_df[feature_columns]

        # Keep a copy of raw data for display
        raw_data = testData_df.values
        
        # Preprocessing using saved label encoders
        for le_entry in label_encoder:
            col_name = le_entry[0]
            le = le_entry[1]
            if col_name in testData_df.columns:
                # Handle unseen labels by assigning -1 or similar
                testData_df[col_name] = testData_df[col_name].astype(str).map(
                    lambda s: le.transform([s])[0] if s in le.classes_ else -1
                )
        
        testData_df.fillna(0, inplace=True)
        
        # Scale data
        scaled_test = scaler.transform(testData_df.values)
        
        # Prediction
        preds = rf_model.predict(scaled_test)
        
        # Prepare data for template rendering via macros
        predictions = [labels[p] for p in preds]
        
        # PERSISTENCE: Store results in session for session-level caching
        session['last_predictions'] = predictions
        session['last_raw_data'] = raw_data.tolist()
        
        # Convert feature columns to list if it exists for JSON serialization
        f_cols = feature_columns.tolist() if hasattr(feature_columns, 'tolist') else feature_columns
        session['last_feature_columns'] = f_cols
        
        return render_template('UserScreen.html', 
                             predictions=predictions, 
                             raw_data=raw_data.tolist(), 
                             page_type='result', 
                             feature_columns=f_cols)
        
    except Exception as e:
        print(f"Error during prediction: {e}")
        flash(f"Analysis failed: {str(e)}", "error")
        return redirect(url_for('predictView'))

def get_genai_insight(attack_type):
    """Simulates a Generative AI explanation for the predicted attack type."""
    insights = {
        "normal": "Traffic patterns match authorized user behavior. No immediate action required.",
        "neptune": "High-volume SYN flood detected. This mimics a 'Neptune' DoS attack. Recommendation: Enable SYN cookies and rate limiting.",
        "back": "Possible buffer overflow attempt on HTTP requests detected. Inspect web server logs for malicious GET requests.",
        "teardrop": "Malformed IP fragments identified. This signature is often used to crash older operating systems. Patch network stack if vulnerable.",
        "satan": "Port scanning activity detected. An attacker is probing your infrastructure for open ports. Consider blacklisting the source IP.",
        "ipsweep": "Network reconnaissance detected. The source is scanning for active hosts in your subnet.",
        "warezclient": "Unauthorized file transfer activity identified. This could indicate data exfiltration or unauthorized software distribution."
    }
    return insights.get(attack_type.lower(), "AI suggests this is a known signature. Monitor for unusual lateral movement in the network.")

# Global tracking for pulse-aware auto-launch
browser_timer = None

def open_browser():
    """Enterprise-grade browser launcher with Chrome Incognito forcing."""
    global PULSE_DETECTED
    if PULSE_DETECTED:
        print("[Smart Launch] Aborting: Active pulse detected.")
        return
        
    import subprocess
    url = "http://127.0.0.1:2026/"
    print(f"[Smart Launch] Initializing Incognito Presentation on {url}")
    
    try:
        # Force chrome incognito for a clean, examiner-ready session
        subprocess.Popen(['start', 'chrome', '--incognito', url], shell=True)
    except Exception as e:
        print(f"[-] Smart Launch Error: {e}")
        import webbrowser
        webbrowser.open(url)

def get_project_tree():
    """Neural Discovery: Builds a dynamic, categorical map of the project file system."""
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) # Root
    tree = {
        "Technical Docs": [],
        "Model Systems": [],
        "ML Notebooks": [],
        "Environment": []
    }
    
    # Exclude list for the explorer (as requested)
    forbidden_dirs = {'.venv', '.git', '.vscode', '__pycache__', 'Software', 'Project materials', '.ipynb_checkpoints'}
    forbidden_files = {'users.json', '.env', '.gitignore'}

    # Strategy: Scan root and CyberAttackPrediction
    scan_targets = [base_dir, os.path.join(base_dir, 'CyberAttackPrediction'), os.path.join(base_dir, 'docs')]
    
    seen_paths = set()

    for target in scan_targets:
        if not os.path.exists(target): continue
        for item in sorted(os.listdir(target)):
            if item in forbidden_dirs or item in forbidden_files or item.startswith('.'):
                continue
            
            full_item_path = os.path.join(target, item)
            if os.path.isdir(full_item_path): continue # sidebar handles file navigation; directory browsing is inline
            
            # Relative path from Root for URIs
            rel_path = os.path.relpath(full_item_path, base_dir).replace('\\', '/')
            if rel_path in seen_paths: continue
            seen_paths.add(rel_path)

            ext = os.path.splitext(item)[1].lower()
            
            # Categorization Logic
            if ext == '.md' or (ext == '.txt' and 'guide' in item.lower()):
                tree["Technical Docs"].append({"name": item, "path": rel_path})
            elif ext == '.docx':
                tree["Technical Docs"].append({"name": item, "path": rel_path})
            elif ext == '.py':
                tree["Model Systems"].append({"name": item, "path": rel_path})
            elif ext == '.ipynb':
                tree["ML Notebooks"].append({"name": item, "path": rel_path})
            elif ext in ['.txt', '.json', '.bat', '.ps1']:
                tree["Environment"].append({"name": item, "path": rel_path})

    return tree

@app.route('/documentation')
def documentation():
    if 'user' not in session:
        return redirect(url_for('UserLogin'))
    return render_template('project overview.html')

@app.route('/explorer')
def explorer():
    if 'user' not in session:
        return redirect(url_for('UserLogin'))
        
    doc_path = request.args.get('doc', 'README.md')
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) # Root
    
    # Neural Discovery Tree
    project_tree = get_project_tree()

    # Hardened security check to prevent directory traversal
    try:
        # Resolve the absolute path and ensure it stays within the restricted project tree
        requested_abs = os.path.abspath(os.path.join(base_dir, doc_path))
        project_root = os.path.abspath(base_dir)
        
        if not requested_abs.startswith(project_root):
            return render_template('documentation.html', error="Security Alert: Outbound traversal blocked.", project_tree=project_tree)
        
        # Explicit check for sensitive files even within project
        forbidden_docs = {'users.json', '.env', 'saved_creds.json'}
        if os.path.basename(requested_abs) in forbidden_docs:
             return render_template('documentation.html', error="Access Restricted: Protected System Resource", project_tree=project_tree)

        full_path = requested_abs
    except Exception:
        return render_template('documentation.html', error="Invalid neural path sequence", project_tree=project_tree)

    if not os.path.exists(full_path):
        return render_template('documentation.html', error=f"System Insight Unavailable: {doc_path}", project_tree=project_tree)
        
    try:
        # FOLDER LISTING INTELLIGENCE
        if os.path.isdir(full_path):
            items = []
            for item in sorted(os.listdir(full_path)):
                if not item.startswith('.') and item not in {'.venv', 'Software', 'Project materials'}:
                    item_rel = os.path.join(doc_path, item).replace('\\', '/')
                    is_dir = os.path.isdir(os.path.join(full_path, item))
                    items.append({'name': item, 'path': item_rel, 'is_dir': is_dir})
            
            render_data = {
                'items': items,
                'current_doc': doc_path.replace('\\', '/'),
                'doc_type': 'folder',
                'project_tree': project_tree
            }
            
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                # SPREAD PURE SYNC: Render only the internal macro from documentation.html
                html = render_template_string(
                    "{% from 'documentation.html' import render_doc_fragment %}{{ render_doc_fragment(doc_type, current_doc, items=items) }}",
                    **render_data
                )
                return jsonify({
                    'html': html,
                    'doc_type': 'folder',
                    'path': render_data['current_doc'],
                    'title': doc_path.split('/')[-1] or 'Repository',
                    'breadcrumb': render_data['current_doc'].replace('CyberAttackPrediction/', '')
                })

            return render_template('documentation.html', **render_data)

        ext = os.path.splitext(doc_path)[1].lower()
        if ext == '.docx':
            content = "[Binary Word Data]" # Placeholder for reading branch
        else:
            with open(full_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
        display_path = doc_path.replace('\\', '/')
        if ext == '.md':
            html_content = markdown2.markdown(content, extras=["tables", "fenced-code-blocks", "toc", "header-ids"])
            # Fix image src paths: markdown/HTML files use project-relative paths
            # e.g. "CyberAttackPrediction/static/images/x.png" or "static/images/x.png"
            # Flask serves static at root (static_url_path=''), so correct URL is "/images/x"
            import re as _re
            html_content = _re.sub(
                r'(src=["\'])(?:(?:\.\./)*)(?:CyberAttackPrediction/)?static/images/',
                r'\1/images/',
                html_content
            )
            doc_type = 'markdown'
        elif ext == '.docx':
            try:
                doc = Document(full_path)
                html = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        # Simple Mapping
                        style = para.style.name.lower()
                        if 'heading 1' in style: html.append(f'<h1>{para.text}</h1>')
                        elif 'heading 2' in style: html.append(f'<h2>{para.text}</h2>')
                        else:
                            # Basic formatting scan
                            p_html = ""
                            for run in para.runs:
                                r_text = run.text
                                if run.bold: r_text = f'<strong>{r_text}</strong>'
                                if run.italic: r_text = f'<em>{r_text}</em>'
                                p_html += r_text
                            html.append(f'<p>{p_html}</p>')
                
                for table in doc.tables:
                    html.append('<div class="table-responsive"><table class="table table-bordered border-glass-soft">')
                    for row in table.rows:
                        html.append('<tr>')
                        for cell in row.cells:
                            html.append(f'<td class="p-2">{cell.text}</td>')
                        html.append('</tr>')
                    html.append('</table></div>')
                
                html_content = "\n".join(html)
                doc_type = 'word'
            except Exception as e:
                html_content = f'<div class="alert alert-warning">Error processing Word document: {str(e)}</div>'
                doc_type = 'word'
        elif ext in ['.py', '.bat', '.ps1', '.json', '.html', '.css', '.js', '.txt']:
            lexer = get_lexer_for_filename(full_path) if ext != '.txt' else TextLexer()
            formatter = HtmlFormatter(style='monokai', full=False, cssclass="highlight")
            html_content = highlight(content, lexer, formatter)
            doc_type = 'code'
            html_content = f'<style>{formatter.get_style_defs(".highlight")}</style>{html_content}'
        else:
            html_content = f'<pre style="white-space: pre-wrap; word-wrap: break-word;">{content}</pre>'
            doc_type = 'text'

        render_data = {
            'md_content': html_content,
            'raw_content': content if ext != '.docx' else "", # DOCX editing will use HTML-to-Text conversion or be read-only for now
            'current_doc': display_path,
            'doc_type': doc_type,
            'project_tree': project_tree
        }

        # AJAX NEURAL SYNC: Return macro-rendered fragment if requested via AJAX
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            # SPREAD PURE SYNC: Render only the internal macro from documentation.html
            html = render_template_string(
                "{% from 'documentation.html' import render_doc_fragment %}{{ render_doc_fragment(doc_type, current_doc, md_content=md_content) }}",
                **render_data
            )
            return jsonify({
                'html': html,
                'doc_type': doc_type,
                'path': display_path,
                'title': display_path.split('/')[-1],
                'breadcrumb': display_path.replace('CyberAttackPrediction/', ''),
                'project_tree': project_tree
            })

        return render_template('documentation.html', **render_data)
        
    except Exception as e:
        return render_template('documentation.html', error=str(e), project_tree=project_tree)

@app.route('/api/explorer/save', methods=['POST'])
def explorer_save():
    if 'user' not in session:
        return jsonify({"status": "error", "message": "Auth Required"}), 401
    
    data = request.json
    rel_path = data.get('path', '')
    new_content = data.get('content', '')
    
    # NEURAL FIREWALL: Absolute path enforcement
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    project_root = os.path.abspath(base_dir)
    requested_abs = os.path.abspath(os.path.join(base_dir, rel_path))
    
    if not requested_abs.startswith(project_root) or any(f in rel_path for f in ['users.json', '.env', 'saved_creds.json']):
        return jsonify({"status": "error", "message": "Neural Integrity Violation: Access Denied"}), 403
        
    full_path = requested_abs
    
    try:
        ext = os.path.splitext(rel_path)[1].lower()
        if ext == '.docx':
            # Save HTML-to-Docx (Simple Text replacement)
            doc = Document()
            # Basic conversion: Split by newlines or use a simple parser
            # For now, we save as plain paragraphs to ensure structure
            for line in new_content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            doc.save(full_path)
        else:
            with open(full_path, 'w', encoding='utf-8') as f:
                f.write(new_content)
        
        return jsonify({"status": "success", "message": "Neural Synchronized: System Hardened"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/explorer/run', methods=['POST'])
def explorer_run():
    if 'user' not in session:
        return jsonify({"status": "error", "message": "Auth Required"}), 401
    
    # Restrict to admin/ultimate users for security
    # Restrict to admin/ultimate users for security
    if not session.get('is_ultimate') and session.get('user') != ADMIN_ID:
         return jsonify({"status": "error", "message": "Neural Access Denied: Administrator Level Restricted"}), 403

    data = request.json
    rel_path = data.get('path', '')
    
    # NEURAL FIREWALL: Prevent execution of the main server script or dangerous utilities via explorer
    restricted_scripts = {'Main.py', 'app.py', 'server.py', '.env', 'users.json', 'saved_creds.json'}
    if os.path.basename(rel_path) in restricted_scripts or '..' in rel_path:
        return jsonify({"status": "error", "message": "Neural Integrity Violation: Execution restricted"}), 403

    if not rel_path or not rel_path.endswith('.py'):
        return jsonify({"status": "error", "message": "Invalid Script Path: Only .py files allowed"}), 400
        
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    full_path = os.path.join(base_dir, rel_path)
    
    if not os.path.exists(full_path):
        return jsonify({"status": "error", "message": "Script not found on disk"}), 404

    try:
        # Determine the best python interpreter (current venv preferred)
        python_exe = sys.executable
        
        # Security: Execute with a timeout to prevent hanging the web server
        process = subprocess.run(
            [python_exe, full_path],
            capture_output=True,
            text=True,
            timeout=30,
            cwd=base_dir # Run from project root
        )
        
        return jsonify({
            "status": "success",
            "stdout": process.stdout,
            "stderr": process.stderr,
            "returncode": process.returncode,
            "message": "Execution Complete" if process.returncode == 0 else "Execution Failed"
        })
        
    except subprocess.TimeoutExpired:
        return jsonify({"status": "error", "message": "Execution Timed Out (30s Limit)"}), 408
    except Exception as e:
        return jsonify({"status": "error", "message": f"Execution Error: {str(e)}"}), 500

@app.route('/Train')
def train_view():
    if 'user' not in session:
        return redirect(url_for('UserLogin'))
    
    # Check if a model has already been trained using the standardized path
    base_dir = os.path.dirname(os.path.abspath(__file__))
    model_path = os.path.join(base_dir, "model", "trained_rf_model.pkl")
    model_ready = os.path.exists(model_path)
    
    train_result = session.get('train_result')
    return render_template('Train.html', page_type='train', model_ready=model_ready, train_result=train_result)

@app.route('/DownloadSample')
def DownloadSample():
    if 'user' not in session:
        return redirect(url_for('index'))
    
    dataset_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Dataset")
    return send_from_directory(dataset_dir, "sample_train.csv", as_attachment=True)


@app.route('/TrainAction', methods=['POST'])
def TrainAction():
    if 'user' not in session:
        return jsonify({"status": "error", "message": "Unauthorized"}), 401
    
    # Check if a custom file was uploaded
    custom_path = None
    if 'training_data' in request.files:
        file = request.files['training_data']
        if file.filename != '':
            # Ensure Dataset directory exists
            base_dir = os.path.dirname(os.path.abspath(__file__))
            dataset_dir = os.path.join(base_dir, "Dataset")
            if not os.path.exists(dataset_dir):
                os.makedirs(dataset_dir)
            
            custom_path = os.path.join(dataset_dir, "custom_train.csv")
            try:
                file.save(custom_path)
            except Exception as e:
                return jsonify({"status": "error", "message": f"Failed to save training file: {str(e)}"})
    
    # Call training with optional path
    global system_status
    system_status = "busy"
    try:
        result = run_training(dataset_path=custom_path)
        
        # If training was successful, reload the model in memory and persist in session
        if result.get('status') == 'success':
            load_ml_model()
            session['train_result'] = result
            
        return jsonify(result)
    finally:
        system_status = "ready"

@app.errorhandler(404)
def handle_404(e):
    """Premium 404 error handler for CyberShield."""
    return render_template('404.html'), 404

@app.errorhandler(500)
def handle_500(e):
    """Secure 500 error handler to prevent internal logic leakage."""
    return render_template('404.html', message="System Integration Error: The request could not be processed safely."), 500

# --- Legal & Contact Routes ---
@app.route('/Contact')
def contact():
    return render_template('Legal.html', section='contact')

@app.route('/TermsOfService')
def terms():
    return render_template('Legal.html', section='terms')

@app.route('/PrivacyPolicy')
def privacy():
    return render_template('Legal.html', section='privacy')

@app.route('/favicon.ico')
def favicon():
    return send_from_directory(os.path.join(app.root_path, 'static', 'images'),
                             'favicon.svg', mimetype='image/svg+xml')

if __name__ == '__main__':
    import sys
    
    # Heartbeat-Aware Smart Launch:
    if '--no-browser' not in sys.argv:
        if os.environ.get("WERKZEUG_RUN_MAIN") == "true" or not app.debug:
            print("[Smart Launch] Starting pulse detector (6.5s)...")
            browser_timer = Timer(6.5, open_browser)
            browser_timer.start()
    
    # use_reloader=False prevents Werkzeug from spawning a second child process.
    # Without this, the reloader restarts the server on every file save, which:
    #   1. Resets PULSE_DETECTED to False in the new process → fires a second browser window
    #   2. Causes WinError 10038 (invalid socket) on Python 3.13 / Windows
    # The launcher handles process management, so the reloader is not needed here.
    app.run(debug=True, port=2026, use_reloader=False)
